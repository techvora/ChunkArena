import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import subprocess
import os
import json
import re

# -----------------------------
# CONFIG
# -----------------------------
TOPICS = [
    "Monetary_policy",
    "Fiscal_policy",
    "Inflation",
    "Interest_rates",
    "Central_banks",
    "Quantitative_easing",
    "Recession",
    "GDP",
    "Banking_system",
    "Financial_crisis"
]

BASE_URL = "https://en.wikipedia.org/wiki/"
OUTPUT_JSON = "finance_wiki.json"
OUTPUT_DOCX = "finance_wiki.docx"

# -----------------------------
# FETCH HTML
# -----------------------------
def fetch_html(topic):
    url = BASE_URL + topic
    headers = {"User-Agent": "Mozilla/5.0"}

    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        print(f"❌ Failed: {topic}")
        return None, url

    return BeautifulSoup(response.text, "lxml"), url


# -----------------------------
# GET CONTENT ROOT (ROBUST)
# -----------------------------
def get_content_root(soup):
    selectors = [
        {"name": "div", "attrs": {"class": "mw-parser-output"}},
        {"name": "div", "attrs": {"id": "mw-content-text"}},
        {"name": "main"}
    ]

    for sel in selectors:
        content = soup.find(sel["name"], sel.get("attrs", {}))
        if content:
            return content

    return None


# -----------------------------
# CLEAN TEXT
# -----------------------------
def clean_text(text):
    text = re.sub(r"\[\d+\]", "", text)  # remove [1]
    text = re.sub(r"\s+", " ", text)
    return text.strip()


# -----------------------------
# PARSE CONTENT
# -----------------------------
def parse_content(soup, topic, url):
    content_div = get_content_root(soup)

    if content_div is None:
        print(f"⚠️ No content found for {topic}")
        return []

    parsed = []
    current_section = "Introduction"

    for el in content_div.descendants:

        # ------------------ HEADINGS ------------------
        if el.name == "h2":
            title = clean_text(el.get_text())
            if title.lower() not in ["references", "external links", "see also"]:
                current_section = title
                parsed.append({
                    "topic": topic,
                    "type": "heading",
                    "section": current_section,
                    "content": title,
                    "source": url
                })

        elif el.name == "h3":
            subtitle = clean_text(el.get_text())
            current_section = subtitle
            parsed.append({
                "topic": topic,
                "type": "subheading",
                "section": current_section,
                "content": subtitle,
                "source": url
            })

        # ------------------ PARAGRAPHS ------------------
        elif el.name == "p":
            text = clean_text(el.get_text())
            if len(text) > 50:
                parsed.append({
                    "topic": topic,
                    "type": "text",
                    "section": current_section,
                    "content": text,
                    "source": url
                })

        # ------------------ TABLES ------------------
        # elif el.name == "table":
        #     rows = []
        #     for tr in el.find_all("tr"):
        #         cols = [clean_text(td.get_text()) for td in tr.find_all(["td", "th"])]
        #         if cols:
        #             rows.append(cols)

        #     if rows:
        #         parsed.append({
        #             "topic": topic,
        #             "type": "table",
        #             "section": current_section,
        #             "content": rows,
        #             "source": url
        #         })

        # ------------------ IMAGES ------------------
        elif el.name == "img":
            src = el.get("src")
            if src and "thumb" in src:
                img_url = "https:" + src if src.startswith("//") else src
                parsed.append({
                    "topic": topic,
                    "type": "image",
                    "section": current_section,
                    "content": img_url,
                    "source": url
                })

        # ------------------ FORMULAS ------------------
        elif el.name == "span" and "mwe-math-element" in el.get("class", []):
            img = el.find("img")
            if img and img.get("alt"):
                parsed.append({
                    "topic": topic,
                    "type": "formula",
                    "section": current_section,
                    "content": img["alt"],
                    "source": url
                })

    return parsed


# -----------------------------
# SAVE JSON
# -----------------------------
def save_json(data, topic):
    filename = f"{topic}.json"

    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print(f"✅ JSON saved: {filename}")


# -----------------------------
# SAVE DOCX
# -----------------------------
def save_docx(data, topic):
    from docx import Document

    filename = f"{topic}.docx"
    doc = Document()

    doc.add_heading(topic.replace("_", " "), 0)

    for item in data:

        if item["type"] == "heading":
            doc.add_heading(item["content"], level=1)

        elif item["type"] == "subheading":
            doc.add_heading(item["content"], level=2)

        elif item["type"] == "text":
            doc.add_paragraph(item["content"])

        elif item["type"] == "table":
            table_data = item["content"]

            if not table_data:
                continue

            max_cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=len(table_data), cols=max_cols)

            for i, row in enumerate(table_data):
                for j in range(max_cols):
                    value = row[j] if j < len(row) else ""
                    table.rows[i].cells[j].text = value

        elif item["type"] == "image":
            doc.add_paragraph(f"Image: {item['content']}")

        elif item["type"] == "formula":
            doc.add_paragraph(f"Formula: {item['content']}")

    doc.save(filename)
    print(f"✅ DOCX saved: {filename}")
    return filename


def convert_docx_to_pdf(docx_file):
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_file,
            "--outdir", os.path.dirname(docx_file)
        ], check=True)

        print(f"✅ PDF generated from {docx_file}")

    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")

# -----------------------------
# MAIN
# -----------------------------
def main():
    for topic in TOPICS:
        print(f"🚀 Processing: {topic}")

        soup, url = fetch_html(topic)
        if soup is None:
            continue

        parsed = parse_content(soup, topic, url)

        # ✅ Save JSON per topic
        # save_json(parsed, topic)

        # ✅ Save DOCX per topic
        docx_file = save_docx(parsed, topic)
        convert_docx_to_pdf(docx_file)


if __name__ == "__main__":
    main()